import streamlit as st
from io import BytesIO
# import gradio as gr
# Def_04 Docx file to translated_Docx file
from transformers import MarianMTModel, MarianTokenizer
import nltk
from nltk.tokenize import sent_tokenize
from nltk.tokenize import LineTokenizer
nltk.download('punkt')
import math
import torch
from docx import Document
from time import sleep
from stqdm import stqdm

import docx
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)
    

if torch.cuda.is_available():  
  dev = "cuda"
else:  
  dev = "cpu" 
device = torch.device(dev)
 
mname = 'Helsinki-NLP/opus-mt-en-hi'
tokenizer = MarianTokenizer.from_pretrained(mname)
model = MarianMTModel.from_pretrained(mname)
model.to(device)

#@st.cache
def btTranslator(docxfile):
  a=getText(docxfile)
  a1=a.split('\n')
  bigtext='''  '''
  for a in a1:
    bigtext=bigtext+'\n'+a
    
  files=Document()
  
  a="/content/drive/MyDrive/Transformers Models/Helsinki-NLP/opus-mt-en-ru"
  b="/content/drive/MyDrive/Transformers Models/Helsinki-NLP/opus-mt-ru-fr"
  c="/content/drive/MyDrive/Transformers Models/Helsinki-NLP/opus-mt-fr-en"
  # d="/content/drive/MyDrive/Transformers Models/Helsinki-NLP/opus-mt-es-en"
  langs=[a,b,c]
  text=bigtext
  
  for _,lang in zip(stqdm(langs),langs):
        sleep(0.5)
        # mname = '/content/drive/MyDrive/Transformers Models/opus-mt-en-hi-Trans Model'
        tokenizer = MarianTokenizer.from_pretrained(lang)
        model = MarianMTModel.from_pretrained(lang)
        model.to(device)
        lt = LineTokenizer()
        batch_size = 8
        paragraphs = lt.tokenize(bigtext)   
        translated_paragraphs = []
        
      for _, paragraph in zip(stqdm(paragraphs),paragraphs):
        # ######################################
          sleep(0.5)

        # ######################################
          sentences = sent_tokenize(paragraph)
          batches = math.ceil(len(sentences) / batch_size)     
          translated = []
          for i in range(batches):
              sent_batch = sentences[i*batch_size:(i+1)*batch_size]
              model_inputs = tokenizer(sent_batch, return_tensors="pt", padding=True, truncation=True, max_length=500).to(device)
              with torch.no_grad():
                  translated_batch = model.generate(**model_inputs)
              translated += translated_batch
          translated = [tokenizer.decode(t, skip_special_tokens=True) for t in translated]
          translated_paragraphs += [" ".join(translated)]
          #files.add_paragraph(translated)
      translated_text = "\n".join(translated_paragraphs)
      bigtext=translated_text
  files.add_paragraph(bigtext) 
  #files=files.save("Translated.docx")
  #binary_output = BytesIO()
  #f=files.save(binary_output)
  #f2=f.getvalue()
  return files


  #return translated_text
st.title('Translator App')
st.markdown("Translate from Docx file")
st.sidebar.subheader("File Upload")

datas=st.sidebar.file_uploader("Original File")
#data=getText("C:\Users\Ambresh C\Desktop\Python Files\Translators\Trail Doc of 500 words.docx")
binary_output = BytesIO()
f3=btTranslator(datas).save(binary_output)
#f4=binary_output(f3)

st.sidebar.download_button(label='Download Translated File',file_name='Translated.docx', data=binary_output.getvalue()) 
# st.text_area(label="",value=btTranslator(datas),height=100)
# Footer
